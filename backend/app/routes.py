from flask import Blueprint, jsonify, request, current_app, send_file
from .models import (
    User, 
    ExerciseSet, 
    Expression, 
    AnswerRecord, 
    ErrorRecord,
    PracticeRecord
)
from . import db
from datetime import datetime, timedelta
from sqlalchemy.exc import IntegrityError
import logging
from random import randint, choice
import random
from .tools.generateExpression import get_expression
from .tools.expression_generator import ExpressionGenerator
# from .tools.generateExpressionController import generateExpressionController
from sqlalchemy import desc
from .config import Config  # 导入配置
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import csv
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from io import StringIO

main = Blueprint('main', __name__)

@main.route('/')
def index():
    return jsonify({"message": "口算练习系统API服务"})

@main.route('/api/test')
def test():
    return jsonify({"message": "API is working!"})

# 用户
@main.route('/api/user', methods=['POST'])
def create_user():
    """创建新用户"""
    current_app.logger.info("Received registration request")
    
    try:
        data = request.get_json()
        current_app.logger.info(f"Received registration data: {data}")
        
        if not data:
            current_app.logger.error("No data received")
            return jsonify({"error": "未接收到注册数据"}), 400
            
        required_fields = ['user_id', 'name', 'id_card', 'grade', 'password']
        missing_fields = [field for field in required_fields if field not in data]
        if missing_fields:
            current_app.logger.error(f"Missing fields: {missing_fields}")
            return jsonify({"error": f"缺少必要的注册信息: {', '.join(missing_fields)}"}), 400

        # 数据验证
        if not isinstance(data['grade'], int):
            current_app.logger.error(f"Invalid grade type: {type(data['grade'])}")
            return jsonify({"error": "年级必须是数字"}), 400
            
        # 先检查用户ID是否已存在
        existing_user = User.query.filter_by(user_id=data['user_id']).first()
        if existing_user:
            return jsonify({"error": f"用户ID '{data['user_id']}' 已被注册"}), 400
            
        # 检查身份证号是否已存在
        existing_id_card = User.query.filter_by(id_card=data['id_card']).first()
        if existing_id_card:
            return jsonify({"error": "该身份证号已被注册"}), 400
            
        # 检查手机号是否已存在（如果提供了手机号）
        if data.get('phone'):
            existing_phone = User.query.filter_by(phone=data['phone']).first()
            if existing_phone:
                return jsonify({"error": "该手机号已被注册"}), 400

        current_app.logger.info("Creating new user")
        new_user = User(
            user_id=data['user_id'],
            name=data['name'],
            id_card=data['id_card'],
            grade=data['grade'],
            phone=data.get('phone')
        )
        new_user.password = data['password']
        
        db.session.add(new_user)
        db.session.commit()
        
        current_app.logger.info("User created successfully")
        return jsonify({"message": "用户注册成功"}), 201
        
    except Exception as e:
        current_app.logger.error(f"Registration error: {str(e)}")
        db.session.rollback()
        return jsonify({"error": str(e)}), 400

@main.route('/api/user/<user_id>', methods=['GET'])
def get_user(user_id):
    """获取用户信息"""
    user = User.query.get(user_id)
    if user:
        return jsonify({
            "user_id": user.user_id,
            "name": user.name,
            "grade": user.grade,
            "id_card": user.id_card,
            "phone": user.phone
        })
    return jsonify({"error": "用户不存在"}), 404

@main.route('/api/user/<user_id>', methods=['PUT'])
def update_user(user_id):
    """更新用户信息"""
    try:
        user = User.query.get(user_id)
        if not user:
            return jsonify({"error": "用户不存在"}), 404
            
        data = request.get_json()
        
        # 更新基本信息
        if 'name' in data:
            user.name = data['name']
        if 'grade' in data:
            user.grade = data['grade']
        if 'phone' in data:
            # 检查手机号是否被其他用户使用
            if data['phone']:
                existing_phone = User.query.filter(
                    User.phone == data['phone'],
                    User.user_id != user_id
                ).first()
                if existing_phone:
                    return jsonify({"error": "该手机号已被其他用户使用"}), 400
            user.phone = data['phone']
            
        # 更新密码
        if 'password' in data:
            user.password = data['password']
            
        db.session.commit()
        return jsonify({"message": "用户信息更新成功"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 400

# 习题集相关接口
@main.route('/api/exercise-set', methods=['POST'])
def create_exercise_set():
    """创建练习集"""
    try:
        data = request.get_json()
        
        # 如果是从错题集生成练习
        if data.get('from_mistakes'):
            # 获取用户的错题
            mistakes = data.get('mistakes', [])
            
            # 创建练习集
            exercise_set = ExerciseSet(
                user_id=data['user_id'],
                total_expressions=len(mistakes),
                bracket_expressions=sum(1 for m in mistakes if '(' in m['expression_text']),
                time_limit=data.get('time_limit', 10),
                operators=','.join([m['expression_text'] for m in mistakes[:20]]),  # 限制题目数量
                operator_count=2,
                min_number=1,
                max_number=100
            )
        else:
            # 正常创建练习集
            exercise_set = ExerciseSet(
                user_id=data['user_id'],
                total_expressions=data['total_expressions'],
                bracket_expressions=data.get('bracket_expressions', 0),
                time_limit=data['time_limit'],
                operators=','.join(data['operators']),
                operator_count=data['operator_count'],
                min_number=data['min_number'],
                max_number=data['max_number']
            )
            
        db.session.add(exercise_set)
        db.session.flush()
        
        # 如果是从错题集生成练习，直接使用已有的题目
        if data.get('from_mistakes'):
            for mistake in mistakes[:20]:  # 限制题目数量
                expression = Expression(
                    exercise_set_id=exercise_set.exercise_set_id,
                    expression_text=mistake['expression_text'],
                    answer=mistake['answer'],
                    has_brackets='(' in mistake['expression_text'],
                    operator_count=sum(1 for c in mistake['expression_text'] if c in '+-×÷')
                )
                db.session.add(expression)
        else:
            # 生成新的题目
            generator = ExpressionGenerator(
                selected_operators=data['operators'],
                operator_count=data['operator_count'],
                min_number=data['min_number'],
                max_number=data['max_number']
            )
            expressions = generator.generate_expressions(
                count=data['total_expressions'],
                bracket_count=data.get('bracket_expressions', 0)
            )
            
            for expr in expressions:
                expression = Expression(
                    exercise_set_id=exercise_set.exercise_set_id,
                    expression_text=expr['expression_text'],
                    answer=expr['answer'],
                    has_brackets='(' in expr['expression_text'],
                    operator_count=sum(1 for c in expr['expression_text'] if c in '+-×÷')
                )
                db.session.add(expression)
        
        db.session.commit()
        
        return jsonify({
            "message": "练习集创建成功",
            "exercise_set_id": exercise_set.exercise_set_id
        })
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating exercise set: {str(e)}")
        return jsonify({"error": str(e)}), 500

@main.route('/api/exercise-set/<int:exercise_set_id>')
def get_exercise_set(exercise_set_id):
    """获取习题集信息"""
    try:
        exercise_set = ExerciseSet.query.get(exercise_set_id)
        if not exercise_set:
            return jsonify({"error": "习题集不存在"}), 404
            
        return jsonify({
            "exercise_set_id": exercise_set.exercise_set_id,
            "total_expressions": exercise_set.total_expressions,
            "bracket_expressions": exercise_set.bracket_expressions,
            "time_limit": exercise_set.time_limit,
            "operators": exercise_set.operators,  # 这里应该已经是字符串格式
            "operator_count": exercise_set.operator_count,
            "min_number": exercise_set.min_number,
            "max_number": exercise_set.max_number,
            "create_time": exercise_set.create_time.isoformat()
        })
        
    except Exception as e:
        current_app.logger.error(f"Error getting exercise set: {str(e)}")
        return jsonify({"error": str(e)}), 500

@main.route('/api/exercise-set/<int:set_id>/expressions', methods=['GET'])
def get_exercise_expressions(set_id):
    """获取习题集的所有表达式"""
    try:
        expressions = Expression.query.filter_by(exercise_set_id=set_id).all()
        return jsonify([{
            'expression_id': expr.expression_id,
            'expression_text': expr.expression_text,
            'answer': expr.answer,
            'has_brackets': expr.has_brackets,
            'operator_count': expr.operator_count
        } for expr in expressions])
    except Exception as e:
        current_app.logger.error(f"Error getting expressions: {str(e)}")
        return jsonify({"error": str(e)}), 400

# 练习记录相关接口
@main.route('/api/practice-record', methods=['POST'])
def create_practice_record():
    """创建练习记录"""
    try:
        data = request.get_json()
        
        # 使用中国时区创建记录
        now = datetime.now(Config.CHINA_TZ)
        
        practice_record = PracticeRecord(
            user_id=data['user_id'],
            exercise_set_id=data['exercise_set_id'],
            duration=data['duration'],
            is_timeout=data['is_timeout'],
            total_expressions=data['total_expressions'],
            bracket_expressions=data['bracket_expressions'],
            time_limit=data['time_limit'],
            operators=data['operators'] if isinstance(data['operators'], str) else ','.join(data['operators']),
            operator_count=data['operator_count'],
            min_number=data['min_number'],
            max_number=data['max_number'],
            completion_time=now  # 使用中国时区的时间
        )
        
        db.session.add(practice_record)
        db.session.commit()
        
        return jsonify({"message": "练习记录创建成功"}), 201
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating practice record: {str(e)}")
        return jsonify({"error": str(e)}), 400

# 获取错题记录
@main.route('/api/error-records/<user_id>')
def get_error_records(user_id):
    try:
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        show_exported = request.args.get('show_exported', 'false').lower() == 'true'
        search = request.args.get('search', '')
        start_date = request.args.get('start_date')
        end_date = request.args.get('end_date')
        
        # 构建基础查询
        query = db.session.query(
            ErrorRecord,
            AnswerRecord,
            Expression,
            PracticeRecord
        ).select_from(ErrorRecord).join(
            AnswerRecord,
            ErrorRecord.answer_record_id == AnswerRecord.answer_record_id
        ).join(
            Expression,
            AnswerRecord.expression_id == Expression.expression_id
        ).join(
            PracticeRecord,
            Expression.exercise_set_id == PracticeRecord.exercise_set_id
        ).filter(
            PracticeRecord.user_id == user_id
        )

        # 根据导出状态过滤
        if not show_exported:
            query = query.filter(ErrorRecord.is_exported == False)
            
        # 搜索过滤
        if search:
            query = query.filter(Expression.expression_text.ilike(f'%{search}%'))
            
        # 日期范围过滤
        if start_date and end_date:
            query = query.filter(
                PracticeRecord.completion_time.between(
                    datetime.strptime(start_date, '%Y-%m-%d'),
                    datetime.strptime(end_date, '%Y-%m-%d') + timedelta(days=1)
                )
            )
        
        # 获取总数
        total = query.count()
        
        # 应用分页
        records = query.order_by(PracticeRecord.completion_time.desc()).offset((page - 1) * per_page).limit(per_page).all()
        
        # 格式化结果
        items = []
        for error, answer, expr, practice in records:
            difficulty = 'simple'
            if expr.has_brackets:
                difficulty = 'hard'
            elif expr.operator_count == 2:
                difficulty = 'medium'
            elif expr.operator_count >= 3:
                difficulty = 'hard'
                
            items.append({
                'id': error.error_record_id,
                'expression': expr.expression_text,
                'correct_answer': expr.answer,
                'user_answer': answer.user_answer,
                'answer_time': answer.answer_time,
                'completion_time': practice.completion_time,
                'difficulty': difficulty,
                'error_count': 1,
                'is_exported': error.is_exported,
                'export_time': error.export_time.isoformat() if error.export_time else None
            })
        
        return jsonify({
            'items': items,
            'total': total,
            'has_exported': False
        })
        
    except Exception as e:
        current_app.logger.error(f"Error getting error records: {str(e)}")
        return jsonify({"error": str(e)}), 500

# 删除错题记录
@main.route('/api/error-records/<error_id>', methods=['DELETE'])
def delete_error_record(error_id):
    try:
        error_record = ErrorRecord.query.get(error_id)
        if not error_record:
            return jsonify({"error": "记录不存在"}), 404
            
        db.session.delete(error_record)
        db.session.commit()
        return jsonify({"message": "删除成功"})
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error deleting error record: {str(e)}")
        return jsonify({"error": str(e)}), 500

# 从错题创建练习
@main.route('/api/practice-from-mistakes', methods=['POST'])
def create_practice_from_mistakes():
    try:
        data = request.get_json()
        user_id = data.get('user_id')
        mistake_ids = data.get('mistake_ids', [])

        # 获取错题对应的原始表达式
        expressions = db.session.query(Expression).join(
            AnswerRecord
        ).join(
            ErrorRecord
        ).filter(
            ErrorRecord.error_record_id.in_(mistake_ids)
        ).all()

        if not expressions:
            return jsonify({"error": "未找到指定的错题"}), 404

        # 创建新的练习集
        exercise_set = ExerciseSet(
            user_id=user_id,
            total_expressions=len(expressions),
            bracket_expressions=sum(1 for e in expressions if e.has_brackets),
            time_limit=30,  # 默认30分钟
            operators=','.join(set(''.join(e.expression_text.split()) for e in expressions)),
            operator_count=max(e.operator_count for e in expressions),
            min_number=0,
            max_number=100
        )
        db.session.add(exercise_set)
        db.session.flush()

        # 复制错题到新练习集
        for expr in expressions:
            new_expr = Expression(
                exercise_set_id=exercise_set.exercise_set_id,
                expression_text=expr.expression_text,
                answer=expr.answer,
                has_brackets=expr.has_brackets,
                operator_count=expr.operator_count
            )
            db.session.add(new_expr)

        db.session.commit()
        return jsonify({
            "message": "练习创建成功",
            "exercise_set_id": exercise_set.exercise_set_id
        }), 201

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating practice from mistakes: {str(e)}")
        return jsonify({"error": str(e)}), 500

@main.route('/api/login', methods=['POST'])
def login():
    """用户登录"""
    data = request.get_json()
    try:
        if not data.get('user_id'):
            return jsonify({"error": "请输入用户ID"}), 400
        if not data.get('password'):
            return jsonify({"error": "请输入密码"}), 400
            
        user = User.query.filter_by(user_id=data['user_id']).first()
        if not user:
            return jsonify({"error": f"用户ID '{data['user_id']}' 不存在"}), 401
        if not user.verify_password(data['password']):
            return jsonify({"error": "密码错误"}), 401
            
        return jsonify({
            "user_id": user.user_id,
            "name": user.name,
            "grade": user.grade
        })
    except Exception as e:
        return jsonify({"error": "登录失败，请稍后重试"}), 400 

@main.route('/api/answer-records/batch', methods=['POST'])
def create_answer_records_batch():
    """批量创建答题记录"""
    try:
        data = request.get_json()
        answers = data.get('answers', [])
        results = []
        
        for answer in answers:
            expression_id = answer.get('expression_id')
            user_answer = answer.get('user_answer')
            answer_time = answer.get('answer_time', 0)
            
            # 获取表达式信息
            expression = Expression.query.get(expression_id)
            if not expression:
                continue
                
            # 判断答案是否正确（允许0.01的误差）
            is_correct = abs(float(user_answer) - expression.answer) < 0.01
            
            # 创建答题记录
            answer_record = AnswerRecord(
                expression_id=expression_id,
                user_answer=user_answer,
                is_correct=is_correct,
                answer_time=answer_time
            )
            db.session.add(answer_record)
            db.session.flush()
            
            # 如果答错了，创建错题记录
            if not is_correct:
                error_record = ErrorRecord(
                    answer_record_id=answer_record.answer_record_id,
                    is_exported=False
                )
                db.session.add(error_record)
            
            results.append({
                "expression_id": expression_id,
                "is_correct": is_correct,
                "correct_answer": expression.answer
            })
        
        db.session.commit()
        return jsonify({"results": results}), 200
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error creating answer records: {str(e)}")
        return jsonify({"error": str(e)}), 500

@main.route('/api/practice-records', methods=['GET'])
def get_practice_records():
    try:
        page = request.args.get('page', 1, type=int)
        size = request.args.get('size', 10, type=int)
        user_id = request.args.get('user_id')
        
        if not user_id:
            return jsonify({"error": "Missing user_id parameter"}), 400
            
        # 获取练习记录
        records = PracticeRecord.query\
            .filter_by(user_id=user_id)\
            .order_by(PracticeRecord.completion_time.desc())\
            .paginate(page=page, per_page=size)
            
        records_data = []
        for record in records.items:
            try:
                # 获取该练习集的所有答题记录
                answer_records = db.session.query(AnswerRecord)\
                    .join(Expression, Expression.expression_id == AnswerRecord.expression_id)\
                    .filter(Expression.exercise_set_id == record.exercise_set_id)\
                    .all()
                    
                # 计算总用时（如果原来的duration为空）
                total_duration = record.duration
                if total_duration is None or total_duration == 0:
                    total_duration = sum(ar.answer_time for ar in answer_records if ar.answer_time is not None)
                
                # 计算正确题数
                correct_count = sum(1 for ar in answer_records if ar.is_correct)
                total_count = len(answer_records)
                
                records_data.append({
                    'id': record.record_id,  # 使用 record_id 而不是 id
                    'exercise_set_id': record.exercise_set_id,
                    'completion_time': record.completion_time.strftime('%Y-%m-%d %H:%M:%S'),
                    'duration': total_duration,  # 使用计算后的总用时
                    'total_expressions': record.total_expressions,
                    'bracket_expressions': record.bracket_expressions,
                    'time_limit': record.time_limit,
                    'operators': record.operators,
                    'operator_count': record.operator_count,
                    'min_number': record.min_number,
                    'max_number': record.max_number,
                    'correct_count': correct_count,
                    'total_count': total_count,
                    'is_timeout': record.is_timeout,
                    'accuracy': f"{(correct_count / total_count * 100):.1f}" if total_count > 0 else "0.0"
                })
            except Exception as e:
                current_app.logger.error(f"Error processing record {record.record_id}: {str(e)}")
                continue
            
        return jsonify({
            'records': records_data,
            'total': records.total
        })
    except Exception as e:
        current_app.logger.error(f"Error getting practice records: {str(e)}")
        return jsonify({'error': str(e)}), 500

@main.route('/api/exercise-set/<int:exercise_set_id>/answers')
def get_exercise_set_answers(exercise_set_id):
    """获取练习集答题记录"""
    try:
        # 获取表达式和答题记录
        expressions = Expression.query.filter_by(exercise_set_id=exercise_set_id).all()
        
        result = []
        for expr in expressions:
            # 获取对应的题记录
            answer_record = AnswerRecord.query.filter_by(expression_id=expr.expression_id).first()
            
            expr_dict = {
                "expression_text": expr.expression_text,
                "answer": expr.answer,
                "user_answer": answer_record.user_answer if answer_record else None,
                "is_correct": answer_record.is_correct if answer_record else False,
                "has_brackets": expr.has_brackets,
                "operator_count": expr.operator_count
            }
            result.append(expr_dict)
        
        # 获取习集信息
        exercise_set = ExerciseSet.query.get(exercise_set_id)
        if not exercise_set:
            return jsonify({"error": "练习集不存在"}), 404
            
        # 返回完整的练习集信息
        return jsonify({
            "exercise_set": {
                "total_expressions": exercise_set.total_expressions,
                "bracket_expressions": exercise_set.bracket_expressions,
                "time_limit": exercise_set.time_limit,
                "operators": exercise_set.operators,
                "operator_count": exercise_set.operator_count,
                "min_number": exercise_set.min_number,
                "max_number": exercise_set.max_number,
                "create_time": exercise_set.create_time.isoformat()
            },
            "expressions": result,
            "statistics": {
                "total": len(expressions),
                "correct": sum(1 for expr in result if expr["is_correct"]),
                "incorrect": sum(1 for expr in result if not expr["is_correct"]),
                "accuracy": f"{(sum(1 for expr in result if expr['is_correct']) / len(expressions) * 100):.1f}%" if expressions else "0%"
            }
        })
        
    except Exception as e:
        current_app.logger.error(f"Error getting exercise set answers: {str(e)}")
        return jsonify({"error": str(e)}), 500 

@main.route('/api/user/<user_id>/statistics')
def get_user_statistics(user_id):
    """获取用户练习统计数据"""
    try:
        # 获取用户所有练习记录
        practice_records = PracticeRecord.query.filter_by(user_id=user_id).all()
        
        # 计算总练习次数（每个练习集只算一次）
        unique_sets = set(record.exercise_set_id for record in practice_records)
        total_practices = len(unique_sets)
        
        # 计算总练习时间（分钟）- 修改时间计算逻辑
        total_minutes = 0
        for record in practice_records:
            # duration 字段存储的是秒数，需要转换为分钟
            if record.duration:
                total_minutes += record.duration / 60
        
        # 获取所有答题记录算正确率
        total_correct = 0
        total_questions = 0
        
        # 使用中国时区的当前时间
        today = datetime.now(Config.CHINA_TZ).date()
        week_stats = []
        
        for i in range(7):
            date = today - timedelta(days=i)
            # 获取当天的所有记录，注意时区转换
            day_records = [r for r in practice_records 
                         if r.completion_time.astimezone(Config.CHINA_TZ).date() == date]
            
            day_correct = 0
            day_total = 0
            day_count = 0
            
            for record in day_records:
                expressions = Expression.query.filter_by(exercise_set_id=record.exercise_set_id).all()
                for expr in expressions:
                    answer_record = AnswerRecord.query.filter_by(expression_id=expr.expression_id).first()
                    if answer_record:
                        day_total += 1
                        day_count += 1
                        if answer_record.is_correct:
                            day_correct += 1
                            total_correct += 1
                        total_questions += 1
            
            day_accuracy = (day_correct / day_total * 100) if day_total > 0 else 0
            
            week_stats.append({
                'date': date.strftime('%m-%d'),
                'accuracy': round(day_accuracy, 1),
                'count': day_count
            })
        
        week_stats.reverse()
        
        # 计算总正确率
        average_accuracy = (total_correct / total_questions * 100) if total_questions > 0 else 0
        
        # 修改答题用分布计算逻辑
        time_distribution = {
            'simple': [0, 0, 0, 0, 0],    # 0-3秒, 3-5秒, 5-10秒, 10-15秒, >15秒
            'medium': [0, 0, 0, 0, 0],
            'hard': [0, 0, 0, 0, 0]
        }
        
        # 获取所有答题记录
        for record in practice_records:
            expressions = Expression.query.filter_by(exercise_set_id=record.exercise_set_id).all()
            for expr in expressions:
                answer_record = AnswerRecord.query.filter_by(expression_id=expr.expression_id).first()
                if answer_record and answer_record.answer_time:
                    # 确定题目难度
                    difficulty = 'simple'
                    if expr.has_brackets:
                        difficulty = 'hard'
                    elif expr.operator_count == 2:
                        difficulty = 'medium'
                    elif expr.operator_count >= 3:
                        difficulty = 'hard'
                    
                    # 使用实际答题时间进行分类
                    time = answer_record.answer_time
                    if time <= 3:
                        time_distribution[difficulty][0] += 1
                    elif time <= 5:
                        time_distribution[difficulty][1] += 1
                    elif time <= 10:
                        time_distribution[difficulty][2] += 1
                    elif time <= 15:
                        time_distribution[difficulty][3] += 1
                    else:
                        time_distribution[difficulty][4] += 1
        
        # 获取运算符统计
        operator_stats = {'+': 0, '-': 0, '×': 0, '÷': 0}
        operator_totals = {'+': 0, '-': 0, '×': 0, '÷': 0}
        
        for record in practice_records:
            expressions = Expression.query.filter_by(exercise_set_id=record.exercise_set_id).all()
            for expr in expressions:
                expr_text = expr.expression_text
                for op in ['×', '÷', '+', '-']:
                    if op in expr_text:
                        operator_totals[op] += 1
                        answer_record = AnswerRecord.query.filter_by(expression_id=expr.expression_id).first()
                        if answer_record and answer_record.is_correct:
                            operator_stats[op] += 1
        
        operator_accuracy = []
        for op in operator_stats:
            if operator_totals[op] > 0:
                accuracy = round(operator_stats[op] / operator_totals[op] * 100, 1)
                operator_accuracy.append({
                    'name': op,
                    'value': accuracy
                })
        
        return jsonify({
            'total_practices': total_practices,
            'total_time': int(total_minutes),
            'average_accuracy': round(average_accuracy, 1),
            'operator_accuracy': operator_accuracy,
            'week_trend': week_stats,
            'time_distribution': time_distribution,
            'time_labels': ['0-3秒', '3-5秒', '5-10秒', '10-15秒', '15秒以上']
        })
        
    except Exception as e:
        current_app.logger.error(f"Error getting user statistics: {str(e)}")
        return jsonify({"error": str(e)}), 500 

@main.route('/api/error-records/<user_id>/export', methods=['POST'])
def export_error_records(user_id):
    try:
        data = request.get_json()
        export_type = data.get('type')
        mistakes = data.get('mistakes', [])
        
        # 更新错题的导出状态
        mistake_ids = [m['id'] for m in mistakes]
        error_records = ErrorRecord.query.filter(
            ErrorRecord.error_record_id.in_(mistake_ids)
        ).all()
        
        for record in error_records:
            record.is_exported = True
            record.export_time = datetime.now(Config.CHINA_TZ)
        
        db.session.commit()
        
        if export_type == 'word':
            doc = Document()
            
            # 设置中文字体
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加标题
            title = doc.add_heading('错题记录', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加基本信息
            doc.add_paragraph(f'导出时间：{datetime.now(Config.CHINA_TZ).strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'错题数量：{len(mistakes)} 道')
            doc.add_paragraph()  # 添加空行
            
            # 添加题目列表
            for idx, mistake in enumerate(mistakes, 1):
                # 添加题编号和算式
                p = doc.add_paragraph()
                p.add_run(f'{idx}. ').bold = True
                p.add_run(f'算式：{mistake["expression"]}')
                
                # 添加答案信息
                details = doc.add_paragraph()
                details.style.font.size = Pt(10)
                # details.add_run('完成时间：').bold = True
                # details.add_run(mistake['completion_time'])
                details.add_run('\n正确答案：').bold = True
                details.add_run(str(mistake['correct_answer']))
                details.add_run('   你的答案：').bold = True
                details.add_run(str(round(mistake['user_answer'],2)))
                # details.add_run('   用时：').bold = True
                # details.add_run(f'{mistake["answer_time"]:.1f}秒')
                details.add_run('   错误次数：').bold = True
                details.add_run(str(mistake['error_count']))
                
                # 添加分隔线
                if idx < len(mistakes):
                    doc.add_paragraph('—' * 50)
            
            # 保存到内存
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            return send_file(
                docx_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'错题本_{datetime.now().strftime("%Y%m%d")}.docx'
            )
            
        else:  # CSV格式
            output = StringIO()
            writer = csv.writer(output)
            
            # 写入表头
            writer.writerow(['序号', '算式', '正确答案', '你的答案', '错误次数'])
            
            # 写入数据
            for idx, mistake in enumerate(mistakes, 1):
                writer.writerow([
                    idx,
                    mistake['expression'],
                    mistake['correct_answer'],
                    round(mistake['user_answer'],2),
                    # f"{mistake['answer_time']:.1f}",
                    # mistake['completion_time'],
                    mistake['error_count']
                ])
            
            # 获取CSV内容并添加BOM
            csv_data = output.getvalue().encode('utf-8-sig')
            
            return send_file(
                BytesIO(csv_data),
                mimetype='text/csv;charset=utf-8',
                as_attachment=True,
                download_name=f'错题本_{datetime.now().strftime("%Y%m%d")}.csv'
            )

    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error exporting error records: {str(e)}")
        return jsonify({"error": str(e)}), 500 

@main.route('/api/exercise-set/<int:exercise_set_id>/complete', methods=['POST'])
def complete_exercise_set(exercise_set_id):
    """完成练习"""
    try:
        data = request.get_json()
        user_id = data.get('user_id')
        duration = data.get('duration', 0)
        is_timeout = bool(data.get('is_timeout', False))  # 确保是布尔值
        
        # 获取练习集信息
        exercise_set = ExerciseSet.query.get(exercise_set_id)
        if not exercise_set:
            return jsonify({"error": "练习集不存在"}), 404
            
        # 创建练习记录
        practice_record = PracticeRecord(
            user_id=str(user_id),  # 确保是字符串
            exercise_set_id=int(exercise_set_id),  # 确保是整数
            completion_time=datetime.now(Config.CHINA_TZ),
            duration=int(duration),  # 确保是整数
            is_timeout=is_timeout,
            is_completed=True,
            # 从练习集获取配置信息
            total_expressions=exercise_set.total_expressions,
            bracket_expressions=exercise_set.bracket_expressions,
            time_limit=exercise_set.time_limit,
            operators=exercise_set.operators,
            operator_count=exercise_set.operator_count,
            min_number=exercise_set.min_number,
            max_number=exercise_set.max_number
        )
        
        db.session.add(practice_record)
        db.session.commit()
        
        return jsonify({"message": "练习完成"}), 200
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error completing exercise set: {str(e)}")
        return jsonify({"error": str(e)}), 500

@main.route('/api/expressions/export', methods=['POST'])
def export_expressions():
    """导出题目"""
    try:
        data = request.get_json()
        format_type = data.get('format', 'csv')
        config = data.get('config')
        
        # 使用表达式生成器生成题目
        generator = ExpressionGenerator(
            selected_operators=config['operators'],
            operator_count=config['operator_count'],
            min_number=config['min_number'],
            max_number=config['max_number']
        )
        expressions = generator.generate_expressions(
            count=config['total_expressions'],
            bracket_count=config.get('bracket_expressions', 0)
        )
        
        if format_type == 'docx':
            # Word格式保持不变，包含答案页
            doc = Document()
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            title = doc.add_heading('口算练习', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            info = doc.add_paragraph()
            info.add_run('练习配置：\n').bold = True
            info.add_run(f'题目数量：{config["total_expressions"]} 道\n')
            if config.get('bracket_expressions', 0) > 0:
                info.add_run(f'括号题目：{config["bracket_expressions"]} 道\n')
            info.add_run(f'运算符：{", ".join(config["operators"])}\n')
            info.add_run(f'数值范围：{config["min_number"]} - {config["max_number"]}\n')
            
            doc.add_paragraph()
            
            # 题目页
            for i in range(0, len(expressions), 2):
                p = doc.add_paragraph()
                p.add_run(f'{i+1}. {expressions[i]["expression_text"]} = ').bold = True
                p.add_run('_' * 8)
                if i + 1 < len(expressions):
                    p.add_run('\t\t')
                    p.add_run(f'{i+2}. {expressions[i+1]["expression_text"]} = ').bold = True
                    p.add_run('_' * 8)
            
            # 答案页
            doc.add_page_break()
            answer_title = doc.add_heading('参考答案', 0)
            answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i in range(0, len(expressions), 4):
                p = doc.add_paragraph()
                for j in range(4):
                    if i + j < len(expressions):
                        p.add_run(f'{i+j+1}. {round(expressions[i+j]["answer"], 2)}\t\t')
            
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            return send_file(
                docx_buffer,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name=f'口算练习_{datetime.now(Config.CHINA_TZ).strftime("%Y%m%d")}.docx'
            )
            
        else:  # CSV 格式
            output = StringIO()
            writer = csv.writer(output)
            
            # 写入表头 - 添加答案列
            writer.writerow(['序号', '算式', '答案'])
            
            # 写入数据 - 答案列留空
            for idx, expr in enumerate(expressions, 1):
                writer.writerow([
                    idx,
                    expr['expression_text'],
                    ''  # 空的答案列供用户填写
                ])
            
            csv_data = output.getvalue().encode('utf-8-sig')
            
            return send_file(
                BytesIO(csv_data),
                mimetype='text/csv;charset=utf-8',
                as_attachment=True,
                download_name=f'口算练习_{datetime.now(Config.CHINA_TZ).strftime("%Y%m%d")}.csv'
            )

    except Exception as e:
        current_app.logger.error(f"Error exporting expressions: {str(e)}")
        return jsonify({"error": str(e)}), 500

@main.route('/api/expressions/import', methods=['POST'])
def import_expressions():
    """导入题目及答案"""
    try:
        if 'file' not in request.files:
            return jsonify({"error": "请选择要上传的文件"}), 400
            
        file = request.files['file']
        user_id = request.form.get('user_id')
        
        if not file or not user_id:
            return jsonify({"error": "缺少必要参数"}), 400
            
        if not file.filename.endswith('.csv'):
            return jsonify({"error": "只支持CSV格式文件"}), 400

        # 读取CSV文件
        content = file.stream.read().decode('utf-8-sig')
        reader = csv.DictReader(StringIO(content))
        
        # 验证CSV文件格式
        required_columns = ['序号', '算式', '答案']
        file_columns = reader.fieldnames
        if not file_columns or not all(col in file_columns for col in required_columns):
            return jsonify({"error": "CSV文件格式错误，需要包含：序号、算式、答案"}), 400
            
        rows = list(reader)
        if not rows:
            return jsonify({"error": "文件中没有题目数据"}), 400

        # 创建练习集
        exercise_set = ExerciseSet(
            user_id=user_id,
            total_expressions=len(rows),
            bracket_expressions=0,
            time_limit=10,
            operators='+,-,×,÷',
            operator_count=2,
            min_number=1,
            max_number=100
        )
        db.session.add(exercise_set)
        db.session.flush()

        # 处理每一行数据
        for row in rows:
            try:
                expression_text = row['算式'].strip()
                user_answer = row['答案'].strip()
                
                if not expression_text:
                    continue
                    
                # 计算正确答案
                correct_answer = eval(expression_text.replace('×', '*').replace('÷', '/'))
                
                # 创建表达式记录
                expression = Expression(
                    exercise_set_id=exercise_set.exercise_set_id,
                    expression_text=expression_text,
                    answer=correct_answer,
                    has_brackets='(' in expression_text,
                    operator_count=sum(1 for c in expression_text if c in '+-×÷')
                )
                db.session.add(expression)
                db.session.flush()
                
                # 如果用户提供了答案，创建答题记录
                if user_answer:
                    try:
                        user_answer_float = float(user_answer)
                        is_correct = abs(user_answer_float - correct_answer) < 0.01
                        
                        answer_record = AnswerRecord(
                            expression_id=expression.expression_id,
                            user_answer=user_answer_float,
                            is_correct=is_correct,
                            answer_time=0
                        )
                        db.session.add(answer_record)
                        db.session.flush()
                        
                        # 如果答错了，创建错题记录
                        if not is_correct:
                            error_record = ErrorRecord(
                                answer_record_id=answer_record.answer_record_id
                            )
                            db.session.add(error_record)
                            
                    except ValueError:
                        continue  # 跳过无效的答案
                    
            except Exception as e:
                current_app.logger.error(f"处理行数据失败: {str(e)}")
                continue  # 跳过处理失败的行
                
        db.session.commit()
        
        return jsonify({
            "message": "导入成功",
            "exercise_set_id": exercise_set.exercise_set_id
        })
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Error importing expressions: {str(e)}")
        return jsonify({"error": f"导入失败: {str(e)}"}), 500